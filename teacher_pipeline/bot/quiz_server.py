"""Веб-API генератора квизов для преподавателей (без Telegram).
Один эндпоинт: браузер шлёт текст ИЛИ файл (+ необязательно страницы/тему) →
DeepSeek генерит квиз → отдаём JSON + meta для дебага."""
import json
import os
import time
import tempfile
from io import BytesIO

import httpx
from docx import Document
from docx.shared import Pt
from fastapi import File, Form, HTTPException, UploadFile
from fastapi import FastAPI
from fastapi.responses import Response
from fastapi.middleware.cors import CORSMiddleware

from pipeline import DEEPSEEK_MODEL, generate_quiz, process_document

GOTENBERG_URL = os.getenv("GOTENBERG_URL", "http://gotenberg:3000")
_LETTERS = ["А", "Б", "В", "Г", "Д", "Е"]


def _plain(text: str) -> str:
    """Убираем сырой LaTeX-разделитель $ для печатного вида."""
    return (text or "").replace("$", "").strip()


def build_quiz_docx(quiz: list, title: str, with_answers: bool) -> bytes:
    """Собирает аккуратный печатный тест (DOCX) в память: шапка-бланк, вопросы,
    варианты буквами, опционально — ключ ответов с новой страницы."""
    doc = Document()
    doc.add_heading(title or "Тест", level=0)
    blank = doc.add_paragraph()
    blank.add_run("ФИО: ______________________    Группа: ____________    Дата: __________").font.size = Pt(11)
    doc.add_paragraph()

    for i, q in enumerate(quiz, 1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. {_plain(q.get('question', ''))}").bold = True
        for j, opt in enumerate(q.get("options", [])):
            letter = _LETTERS[j] if j < len(_LETTERS) else str(j + 1)
            doc.add_paragraph(f"{letter}) {_plain(opt)}")
        doc.add_paragraph()

    if with_answers:
        doc.add_page_break()
        doc.add_heading("Ключ ответов", level=1)
        for i, q in enumerate(quiz, 1):
            ci = q.get("correct", 0)
            letter = _LETTERS[ci] if isinstance(ci, int) and ci < len(_LETTERS) else "?"
            line = f"{i} — {letter}"
            expl = _plain(q.get("explanation", ""))
            if expl:
                line += f"  ({expl})"
            doc.add_paragraph(line)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def docx_to_pdf(docx_bytes: bytes) -> bytes:
    """Конвертирует DOCX → PDF через Gotenberg (LibreOffice), как в сервисе формул."""
    url = f"{GOTENBERG_URL}/forms/libreoffice/convert"
    try:
        resp = httpx.post(url, files={"files": ("quiz.docx", docx_bytes)}, timeout=120.0)
        resp.raise_for_status()
    except Exception as exc:
        raise HTTPException(status_code=502, detail=f"Gotenberg не сконвертировал в PDF: {exc}")
    if not resp.content.startswith(b"%PDF"):
        raise HTTPException(status_code=502, detail="Gotenberg вернул не PDF")
    return resp.content


def parse_pages(spec: str) -> list:
    """«1-10», «1,3,5», «1-3,7» → [1,2,…]. Пусто → None (авто)."""
    spec = (spec or "").strip()
    if not spec:
        return None
    pages: set = set()
    for part in spec.split(","):
        part = part.strip()
        if "-" in part:
            try:
                a, b = part.split("-", 1)
                pages.update(range(int(a), int(b) + 1))
            except ValueError:
                continue
        elif part.isdigit():
            pages.add(int(part))
    return sorted(p for p in pages if p >= 1) or None


app = FastAPI(title="Teacher Quiz Server")
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.get("/health")
def health() -> dict:
    return {"status": "ok", "model": DEEPSEEK_MODEL}


@app.post("/quiz")
async def quiz(
    text: str = Form(""),
    file: UploadFile = File(None),
    pages: str = Form(""),
    topic: str = Form(""),
    num_questions: int = Form(5),
) -> dict:
    """Единая точка генерации квиза. Если приложен файл — квиз по документу
    (с учётом страниц/темы), иначе — по вставленному тексту."""
    num = max(1, min(int(num_questions), 10))
    topic_clean = topic.strip() or None
    t0 = time.perf_counter()

    # Режим «документ»
    if file is not None and file.filename:
        page_range = parse_pages(pages)
        suffix = os.path.splitext(file.filename)[1] or ".txt"
        tmp_path = None
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
                tmp.write(await file.read())
                tmp_path = tmp.name

            result = process_document(tmp_path, num_questions=num, page_range=page_range, topic=topic_clean)
            elapsed_ms = int((time.perf_counter() - t0) * 1000)

            if not result.get("quiz"):
                raise HTTPException(status_code=502, detail="Не удалось сгенерировать квиз (пустой текст/страницы или ошибка DeepSeek)")

            return {
                "quiz": result["quiz"],
                "meta": {
                    "model": DEEPSEEK_MODEL,
                    "source": "document",
                    "filename": file.filename,
                    "pages": pages or "все/авто",
                    "topic": topic or "—",
                    "words": result.get("words"),
                    "generated": len(result["quiz"]),
                    "connectivity": result.get("connectivity"),
                    "elapsed_ms": elapsed_ms,
                },
            }
        finally:
            if tmp_path and os.path.exists(tmp_path):
                os.unlink(tmp_path)

    # Режим «текст»
    if text and text.strip():
        result = generate_quiz(text, num, topic_clean)
        elapsed_ms = int((time.perf_counter() - t0) * 1000)
        if not result:
            raise HTTPException(status_code=502, detail="DeepSeek не вернул квиз (проверьте ключ/баланс)")
        return {
            "quiz": result,
            "meta": {
                "model": DEEPSEEK_MODEL,
                "source": "text",
                "topic": topic or "—",
                "input_chars": len(text),
                "generated": len(result),
                "elapsed_ms": elapsed_ms,
            },
        }

    raise HTTPException(status_code=400, detail="Вставьте текст или выберите файл")


@app.post("/export")
async def export(
    quiz: str = Form(...),
    format: str = Form("docx"),
    title: str = Form("Тест"),
    with_answers: bool = Form(False),
) -> Response:
    """Печатная версия квиза: DOCX (python-docx) или PDF (тот же DOCX через Gotenberg)."""
    try:
        items = json.loads(quiz)
        assert isinstance(items, list) and items
    except Exception:
        raise HTTPException(status_code=400, detail="Пустой или некорректный квиз")

    docx_bytes = build_quiz_docx(items, title.strip() or "Тест", with_answers)

    if format == "pdf":
        pdf = docx_to_pdf(docx_bytes)
        return Response(
            content=pdf,
            media_type="application/pdf",
            headers={"Content-Disposition": 'attachment; filename="quiz.pdf"'},
        )
    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": 'attachment; filename="quiz.docx"'},
    )
